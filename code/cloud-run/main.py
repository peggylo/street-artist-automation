"""
街頭藝人申請系統 - Cloud Run 主程式
Phase 5: 文件處理系統

主要功能：
1. 接收來自 GAS 的申請資料
2. 下載 Word 模板並填寫
3. 轉換為 PDF
4. 上傳到 Google Drive
5. 更新 Google Sheets 狀態
"""

import os
import json
import logging
import tempfile
import subprocess
from datetime import datetime
import pytz
from flask import Flask, request, jsonify
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload, MediaIoBaseDownload
from docx import Document
import io

from config import config

# 設定日誌
logging.basicConfig(
    level=getattr(logging, config.LOGGING["LEVEL"]),
    format=config.LOGGING["FORMAT"]
)
logger = logging.getLogger(__name__)

# 初始化 Flask 應用
app = Flask(__name__)

class DocumentProcessor:
    """文件處理器"""
    
    def __init__(self):
        """初始化 Google API 客戶端"""
        try:
            # 取得服務帳戶憑證
            service_account_info = config.get_service_account_info()
            credentials = service_account.Credentials.from_service_account_info(
                service_account_info,
                scopes=[
                    'https://www.googleapis.com/auth/drive',
                    'https://www.googleapis.com/auth/spreadsheets'
                ]
            )
            
            # 初始化 API 客戶端
            self.drive_service = build('drive', 'v3', credentials=credentials)
            self.sheets_service = build('sheets', 'v4', credentials=credentials)
            
            logger.info("Google API 客戶端初始化成功")
            
        except Exception as e:
            logger.error(f"初始化 Google API 客戶端失敗: {str(e)}")
            raise
    
    def download_template(self, temp_dir):
        """
        從 Google Drive 下載 Word 模板
        
        Args:
            temp_dir (str): 臨時目錄路徑
            
        Returns:
            str: 下載的模板檔案路徑
        """
        try:
            logger.info("開始下載 Word 模板")
            
            # 直接使用模板檔案 ID（方案 B）
            template_file_id = config.GOOGLE_DRIVE["TEMPLATE_WORD_FILE_ID"]
            logger.info(f"使用 Word 模板檔案 ID: {template_file_id}")
            
            # 下載檔案
            request = self.drive_service.files().get_media(fileId=template_file_id)
            file_name = config.GOOGLE_DRIVE["TEMPLATE_FILE_NAME"]
            template_path = os.path.join(temp_dir, file_name)
            
            with open(template_path, 'wb') as f:
                downloader = MediaIoBaseDownload(f, request)
                done = False
                while done is False:
                    status, done = downloader.next_chunk()
            
            logger.info(f"模板下載完成: {template_path}")
            return template_path
            
        except Exception as e:
            logger.error(f"下載模板失敗: {str(e)}")
            raise
    
    def download_copied_file(self, copied_file_id, temp_dir):
        """
        從 Google Drive 下載已複製的 Word 檔案（方案 B）
        
        Args:
            copied_file_id (str): 已複製檔案的 ID
            temp_dir (str): 臨時目錄路徑
            
        Returns:
            str: 下載的檔案路徑
        """
        try:
            logger.info(f"開始下載已複製的 Word 檔案: {copied_file_id}")
            
            # 取得檔案資訊
            file_metadata = self.drive_service.files().get(fileId=copied_file_id).execute()
            file_name = file_metadata.get('name', 'copied_template.docx')
            
            logger.info(f"檔案名稱: {file_name}")
            
            # 下載檔案內容
            request = self.drive_service.files().get_media(fileId=copied_file_id)
            
            copied_file_path = os.path.join(temp_dir, file_name)
            
            with open(copied_file_path, 'wb') as copied_file:
                downloader = MediaIoBaseDownload(copied_file, request)
                done = False
                while done is False:
                    status, done = downloader.next_chunk()
                    logger.info(f"下載進度: {int(status.progress() * 100)}%")
            
            logger.info(f"已複製檔案下載完成: {copied_file_path}")
            return copied_file_path
            
        except Exception as e:
            logger.error(f"下載已複製檔案失敗: {str(e)}")
            raise
    
    def fill_template(self, template_path, application_data, output_path):
        """
        填寫 Word 模板
        
        Args:
            template_path (str): 模板檔案路徑
            application_data (dict): 申請資料
            output_path (str): 輸出檔案路徑
        """
        try:
            logger.info("開始填寫 Word 模板")
            
            # 開啟 Word 文件
            doc = Document(template_path)
            
            # 準備替換資料
            replacements = {
                config.TEMPLATE_PROCESSING["URL_PLACEHOLDER"]: application_data.get("video_url", ""),
            }
            
            # 處理日期替換
            dates = application_data.get("selected_dates", [])
            for i, placeholder in enumerate(config.TEMPLATE_PROCESSING["DATE_PLACEHOLDERS"]):
                if i < len(dates):
                    # 格式化日期為 YYYY/MM/DD 格式
                    date_obj = dates[i]
                    if isinstance(date_obj, dict):
                        # 如果是字典，提取 display 欄位
                        date_str = date_obj.get("display", "")
                    else:
                        # 如果是字串，直接使用
                        date_str = str(date_obj)
                    replacements[placeholder] = date_str
                else:
                    # 空白處理
                    replacements[placeholder] = ""
            
            logger.info(f"替換資料: {replacements}")
            
            # 執行文字替換
            for paragraph in doc.paragraphs:
                for placeholder, value in replacements.items():
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, value)
                        logger.info(f"替換 {placeholder} -> {value}")
            
            # 處理表格中的文字
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for placeholder, value in replacements.items():
                            if placeholder in cell.text:
                                cell.text = cell.text.replace(placeholder, value)
                                logger.info(f"表格中替換 {placeholder} -> {value}")
            
            # 儲存填寫後的文件
            doc.save(output_path)
            logger.info(f"Word 模板填寫完成: {output_path}")
            
        except Exception as e:
            logger.error(f"填寫模板失敗: {str(e)}")
            raise
    
    def convert_to_pdf(self, word_path, temp_dir):
        """
        使用 LibreOffice 將 Word 轉換為 PDF
        
        Args:
            word_path (str): Word 檔案路徑
            temp_dir (str): 臨時目錄路徑
            
        Returns:
            str: PDF 檔案路徑
        """
        try:
            logger.info("開始轉換 PDF")
            
            # 建構 LibreOffice 指令
            cmd = [
                config.LIBREOFFICE["COMMAND"],
                "--headless",
                "--convert-to", "pdf",
                "--outdir", temp_dir,
                word_path
            ]
            
            # 執行轉換
            result = subprocess.run(
                cmd,
                timeout=config.LIBREOFFICE["TIMEOUT_SECONDS"],
                capture_output=True,
                text=True
            )
            
            if result.returncode != 0:
                raise Exception(f"LibreOffice 轉換失敗: {result.stderr}")
            
            # 找到生成的 PDF 檔案
            word_filename = os.path.basename(word_path)
            pdf_filename = os.path.splitext(word_filename)[0] + ".pdf"
            pdf_path = os.path.join(temp_dir, pdf_filename)
            
            if not os.path.exists(pdf_path):
                raise Exception(f"找不到轉換後的 PDF: {pdf_path}")
            
            logger.info(f"PDF 轉換完成: {pdf_path}")
            return pdf_path
            
        except Exception as e:
            logger.error(f"PDF 轉換失敗: {str(e)}")
            raise
    
    def upload_word(self, word_path, application_data):
        """
        上傳 Word 到 Google Drive（方案 B：覆蓋現有檔案）
        
        Args:
            word_path (str): Word 檔案路徑
            application_data (dict): 申請資料
            
        Returns:
            str: 上傳後的檔案連結
        """
        try:
            # 檢查是否為方案 B（有 copiedFileId）
            copied_file_id = application_data.get("copiedFileId")
            
            if copied_file_id:
                logger.info(f"方案 B: 覆蓋現有 Word 檔案 {copied_file_id}")
                
                # 覆蓋現有檔案
                media = MediaFileUpload(
                    word_path, 
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                
                file = self.drive_service.files().update(
                    fileId=copied_file_id,
                    media_body=media,
                    fields='id,webViewLink'
                ).execute()
                
                file_id = file.get('id')
                file_url = file.get('webViewLink')
                
                logger.info(f"Word 覆蓋完成: {file_url}")
                return file_url
                
            else:
                logger.info("方案 A: 建立新 Word 檔案")
                
                # 生成檔案名稱
                year = application_data.get("year")
                month = application_data.get("month")
                
                # 檢查必要參數
                if not year or not month:
                    raise ValueError(f"缺少必要參數: year={year}, month={month}")
                
                # 使用與 PDF 類似的命名規則，但加上 _待處理 後綴
                timestamp = application_data.get("timestamp", "")
                filename = f"申請表_{year}年{month}月_{timestamp}_待處理.docx"
                
                # 建立新檔案
                folder_id = config.GOOGLE_DRIVE["GENERATED_FOLDER_ID"]
                media = MediaFileUpload(
                    word_path, 
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                
                file_metadata = {
                    'name': filename,
                    'parents': [folder_id]
                }
                
                file = self.drive_service.files().create(
                    body=file_metadata,
                    media_body=media,
                    fields='id,webViewLink'
                ).execute()
                
                file_id = file.get('id')
                file_url = file.get('webViewLink')
                
                logger.info(f"Word 上傳完成: {file_url}")
                return file_url
            
        except Exception as e:
            logger.error(f"Word 上傳失敗: {str(e)}")
            raise
    
    def upload_pdf(self, pdf_path, application_data):
        """
        上傳 PDF 到 Google Drive（方案 B：覆蓋現有檔案）
        
        Args:
            pdf_path (str): PDF 檔案路徑
            application_data (dict): 申請資料
            
        Returns:
            str: 上傳後的檔案連結
        """
        try:
            # 檢查是否為方案 B（有 pdfFileId）
            pdf_file_id = application_data.get("pdfFileId")
            
            if pdf_file_id:
                logger.info(f"方案 B: 覆蓋現有 PDF 檔案 {pdf_file_id}")
                
                # 覆蓋現有檔案
                media = MediaFileUpload(pdf_path, mimetype='application/pdf')
                
                file = self.drive_service.files().update(
                    fileId=pdf_file_id,
                    media_body=media,
                    fields='id,webViewLink'
                ).execute()
                
                file_id = file.get('id')
                file_url = file.get('webViewLink')
                
                logger.info(f"PDF 覆蓋完成: {file_url}")
                return file_url
                
            else:
                logger.info("方案 A: 建立新 PDF 檔案")
                
                # 生成檔案名稱
                year = application_data.get("year")
                month = application_data.get("month")
                
                # 檢查必要參數
                if not year or not month:
                    raise ValueError(f"缺少必要參數: year={year}, month={month}")
                
                filename = config.generate_pdf_filename(year, month)
                
                # 建立新檔案
                folder_id = config.GOOGLE_DRIVE["GENERATED_FOLDER_ID"]
                media = MediaFileUpload(pdf_path, mimetype='application/pdf')
                
                file_metadata = {
                    'name': filename,
                    'parents': [folder_id]
                }
                
                file = self.drive_service.files().create(
                    body=file_metadata,
                    media_body=media,
                    fields='id,webViewLink'
                ).execute()
                
                file_id = file.get('id')
                file_url = file.get('webViewLink')
                
                logger.info(f"PDF 上傳完成: {file_url}")
                return file_url
            
        except Exception as e:
            logger.error(f"PDF 上傳失敗: {str(e)}")
            raise
    
    def update_sheets_status(self, user_id, application_data, pdf_url, status="完成", error_message=""):
        """
        更新 Google Sheets 狀態
        
        Args:
            user_id (str): 用戶 ID
            application_data (dict): 申請資料（包含時間戳記）
            pdf_url (str): PDF 檔案連結
            status (str): 狀態
            error_message (str): 錯誤訊息
        """
        try:
            logger.info(f"更新 Sheets 狀態: {status}")
            
            spreadsheet_id = config.GOOGLE_SHEETS["APPLICATION_RECORD_ID"]
            sheet_name = config.GOOGLE_SHEETS["SHEET_NAME"]
            
            # 讀取現有資料，找到對應的用戶記錄
            range_name = f"{sheet_name}!A:K"
            result = self.sheets_service.spreadsheets().values().get(
                spreadsheetId=spreadsheet_id,
                range=range_name
            ).execute()
            
            values = result.get('values', [])
            
            # 改用時間戳記找到精確的記錄
            target_row = None
            target_timestamp = application_data.get("timestamp")
            
            if not target_timestamp:
                # 向後相容：如果沒有時間戳記，回退到原來的邏輯
                logger.warning("沒有時間戳記，使用 User ID 搜尋")
                for i, row in enumerate(values):
                    if len(row) > 1 and row[1] == user_id and len(row) > 6 and row[6] == "待處理":
                        target_row = i + 1
                        break
            else:
                # 用時間戳記精確搜尋
                logger.info(f"使用時間戳記搜尋記錄: {target_timestamp}")
                for i, row in enumerate(values):
                    if len(row) > 0 and row[0] == target_timestamp:
                        target_row = i + 1
                        logger.info(f"找到匹配記錄在第 {target_row} 行")
                        break
            
            if not target_row:
                if target_timestamp:
                    raise Exception(f"找不到時間戳記 {target_timestamp} 的申請記錄")
                else:
                    raise Exception(f"找不到用戶 {user_id} 的待處理記錄")
            
            # 更新狀態 - 使用台灣時區的 YYYYMMDD-HHmmss 統一時間格式
            taiwan_tz = pytz.timezone('Asia/Taipei')
            now_dt = datetime.now(taiwan_tz)
            now = f"{now_dt.year:04d}{now_dt.month:02d}{now_dt.day:02d}-{now_dt.hour:02d}{now_dt.minute:02d}{now_dt.second:02d}"
            update_data = []
            
            if status == "完成":
                # 成功完成：更新狀態、錯誤訊息、PDF路徑、處理開始時間、處理完成時間
                update_data = [
                    [status, "", pdf_url, now, now]  # G, H, I, J, K 欄位
                ]
                update_range = f"{sheet_name}!G{target_row}:K{target_row}"
            else:
                # 處理失敗：更新狀態、錯誤訊息、處理開始時間、處理完成時間
                update_data = [
                    [status, error_message, "", now, now]  # G, H, I, J, K 欄位
                ]
                update_range = f"{sheet_name}!G{target_row}:K{target_row}"
            
            # 執行更新
            self.sheets_service.spreadsheets().values().update(
                spreadsheetId=spreadsheet_id,
                range=update_range,
                valueInputOption='USER_ENTERED',
                body={'values': update_data}
            ).execute()
            
            logger.info(f"Sheets 狀態更新完成: 行 {target_row}")
            
        except Exception as e:
            logger.error(f"更新 Sheets 狀態失敗: {str(e)}")
            raise

# 全域文件處理器實例
doc_processor = DocumentProcessor()

@app.route('/health', methods=['GET'])
def health_check():
    """健康檢查端點"""
    return jsonify({
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "service": "document-processor"
    })

@app.route('/process-application', methods=['POST'])
def process_application():
    """
    Phase 5-6 整合：處理申請文件 + 網站自動化
    
    預期的 JSON 格式：
    {
        "user_id": "用戶ID",
        "timestamp": "20251012-0316",
        "application_data": {
            "year": "2025",
            "month": "10",
            "selected_dates": ["2025/10/5"],
            "video_url": "https://drive.google.com/...",
            "video_source": "常用影片",
            "copiedFileId": "...",
            "pdfFileId": "..."
        },
        "gas_callback_url": "GAS回調URL"  # Phase 6: 新增回調URL
    }
    """
    try:
        # 解析請求資料
        application_data = request.get_json()
        if not application_data:
            return jsonify({"error": "缺少申請資料"}), 400
        
        user_id = application_data.get("user_id")
        timestamp = application_data.get("timestamp")
        gas_callback_url = application_data.get("gas_callback_url")
        
        # 提取申請資料
        app_data = application_data.get("application_data")
        if not app_data:
            return jsonify({"error": "缺少申請資料"}), 400
        
        # 將時間戳記加入申請資料中，供 update_sheets_status 使用
        app_data["timestamp"] = timestamp
        
        logger.info(f"🚀 Phase 5-6 整合流程開始")
        logger.info(f"👤 用戶: {user_id}, 時間戳記: {timestamp}")
        logger.info(f"📋 申請資料: {app_data}")
        
        # ===== Phase 5: 文件處理 =====
        logger.info("📄 Phase 5: 開始文件處理...")
        
        # 更新 Sheets 狀態為「文件處理中」
        doc_processor.update_sheets_status(user_id, app_data, "", "文件處理中")
        
        # 檢查是否為方案 B（GAS 複製 + Cloud Run 編輯）
        copied_file_id = app_data.get("copiedFileId")
        pdf_url = ""
        
        if copied_file_id:
            logger.info(f"使用方案 B: 編輯已複製檔案 {copied_file_id}")
            
            # 建立臨時目錄
            with tempfile.TemporaryDirectory() as temp_dir:
                logger.info(f"使用臨時目錄: {temp_dir}")
                
                # 1. 下載已複製的 Word 檔案
                copied_word_path = doc_processor.download_copied_file(copied_file_id, temp_dir)
                
                # 2. 填寫模板
                filled_word_path = os.path.join(temp_dir, "filled_template.docx")
                doc_processor.fill_template(copied_word_path, app_data, filled_word_path)
                
                # 2.5. 上傳填寫後的 Word 回 Google Drive
                word_url = doc_processor.upload_word(filled_word_path, app_data)
                logger.info(f"Word 檔案已上傳: {word_url}")
                
                # 3. 轉換為 PDF
                pdf_path = doc_processor.convert_to_pdf(filled_word_path, temp_dir)
                
                # 4. 上傳 PDF
                pdf_url = doc_processor.upload_pdf(pdf_path, app_data)
                logger.info(f"PDF 檔案已上傳: {pdf_url}")
        else:
            logger.info("使用方案 A: 下載模板檔案")
            
            # 建立臨時目錄
            with tempfile.TemporaryDirectory() as temp_dir:
                logger.info(f"使用臨時目錄: {temp_dir}")
                
                # 1. 下載模板
                template_path = doc_processor.download_template(temp_dir)
                
                # 2. 填寫模板
                filled_word_path = os.path.join(temp_dir, "filled_template.docx")
                doc_processor.fill_template(template_path, app_data, filled_word_path)
                
                # 2.5. 上傳填寫後的 Word 到 Google Drive
                word_url = doc_processor.upload_word(filled_word_path, app_data)
                logger.info(f"Word 檔案已上傳: {word_url}")
                
                # 3. 轉換為 PDF
                pdf_path = doc_processor.convert_to_pdf(filled_word_path, temp_dir)
                
                # 4. 上傳 PDF
                pdf_url = doc_processor.upload_pdf(pdf_path, app_data)
                logger.info(f"PDF 檔案已上傳: {pdf_url}")
        
        logger.info("✅ Phase 5: 文件處理完成")
        
        # ===== 階段 5: Shortcut 半自動化方案（跳過網站自動化）=====
        logger.info("📱 階段 5: 準備 Shortcut 半自動化方案")
        
        # 更新 Sheets 狀態為「完成」
        doc_processor.update_sheets_status(user_id, app_data, pdf_url, "完成", "")
        logger.info("✅ Sheets 狀態已更新為「完成」")
        
        # 回調 GAS（如果有提供回調 URL）
        if gas_callback_url:
            # 取得 group_id（從請求資料中）
            group_id = application_data.get("group_id")
            
            callback_data = {
                "success": True,
                "user_id": user_id,
                "group_id": group_id,
                "timestamp": timestamp,
                "pdf_file_id": app_data.get("pdfFileId"),
                "message": "✅ 申請表已準備好"
            }
            
            logger.info("📤 準備回調 GAS")
            logger.info(f"📋 回調資料: {callback_data}")
            
            try:
                import requests
                callback_response = requests.post(
                    gas_callback_url,
                    json=callback_data,
                    timeout=10
                )
                logger.info(f"✅ 已回調 GAS: {callback_response.status_code}")
            except Exception as callback_error:
                logger.error(f"⚠️ 回調 GAS 失敗: {str(callback_error)}")
        
        logger.info("🎉 階段 5: 文件處理和回調完成")
        
        return jsonify({
            "success": True,
            "message": "申請處理完成，Shortcut 連結已發送",
            "pdf_url": pdf_url,
            "pdf_file_id": app_data.get("pdfFileId"),
            "user_id": user_id
        })
        
    except Exception as e:
        logger.error(f"❌ 處理申請失敗: {str(e)}")
        
        # 更新失敗狀態
        try:
            user_id = application_data.get("user_id") if application_data else "unknown"
            group_id = application_data.get("group_id") if application_data else None
            app_data = application_data.get("application_data", {}) if application_data else {}
            error_message = f"[文件處理] {str(e)}"
            doc_processor.update_sheets_status(user_id, app_data, "", "失敗", error_message)
            
            # 回調 GAS（失敗通知）
            gas_callback_url = application_data.get("gas_callback_url") if application_data else None
            if gas_callback_url:
                callback_data = {
                    "success": False,
                    "user_id": user_id,
                    "group_id": group_id,
                    "timestamp": application_data.get("timestamp", ""),
                    "message": f"文件處理失敗: {str(e)}"
                }
                
                logger.info(f"📤 準備回調 GAS（失敗通知）")
                logger.info(f"📋 回調資料: {callback_data}")
                
                try:
                    import requests
                    requests.post(gas_callback_url, json=callback_data, timeout=10)
                    logger.info("✅ 已回調 GAS（失敗通知）")
                except Exception as callback_error:
                    logger.error(f"⚠️ 回調 GAS 失敗: {str(callback_error)}")
        except Exception as notify_error:
            logger.error(f"❌ 通知處理失敗: {str(notify_error)}")
        
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

@app.route('/website-automation', methods=['POST'])
def website_automation():
    """
    Phase 6: 網站自動化端點
    
    預期的 JSON 格式：
    {
        "user_id": "用戶ID",
        "application_data": {
            "year": "2025",
            "month": "10",
            "selected_dates": ["2025/10/5"],
            "video_url": "https://drive.google.com/...",
            "video_source": "常用影片",
            "timestamp": "20251012-0316"
        }
    }
    """
    try:
        # 解析請求資料
        request_data = request.get_json()
        if not request_data:
            return jsonify({"error": "缺少請求資料"}), 400
        
        user_id = request_data.get("user_id")
        application_data = request_data.get("application_data", {})
        
        if not user_id:
            return jsonify({"error": "缺少用戶ID"}), 400
        
        logger.info(f"開始網站自動化處理: 用戶 {user_id}")
        
        # 導入並執行網站自動化
        from website_automation_cloud import WebsiteAutomationCloud
        
        # 建立自動化實例（階段 2B：Cloud Run 測試，不提交）
        automation = WebsiteAutomationCloud(stage="2B")
        
        # 執行網站自動化
        result = automation.run_automation(application_data)
        
        if result['success']:
            logger.info(f"網站自動化成功: 用戶 {user_id}")
            return jsonify({
                "success": True,
                "message": "網站自動化完成",
                "result": result,
                "user_id": user_id
            })
        else:
            logger.error(f"網站自動化失敗: 用戶 {user_id}, 錯誤: {result.get('error')}")
            return jsonify({
                "success": False,
                "error": result.get('error', '未知錯誤'),
                "result": result
            }), 500
            
    except Exception as e:
        logger.error(f"網站自動化處理失敗: {str(e)}")
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

if __name__ == '__main__':
    logger.info("啟動文件處理與網站自動化服務")
    app.run(host='0.0.0.0', port=config.HTTP["PORT"])
